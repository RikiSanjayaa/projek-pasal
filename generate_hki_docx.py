from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

output_path = r"C:\Users\msari\.gemini\antigravity\brain\76c8b95b-694e-4f31-b570-9432424eea6a\Deskripsi_HKI_Aplikasi_Pasal.docx"
images_dir = r"e:\projeck git\projek-pasal\extracted_images"

def create_hki_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('DESKRIPSI CIPTAAN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph('JUDUL CIPTAAN\t: Aplikasi Pasal Mobile dan Dashboard Admin')
    doc.add_paragraph('JENIS CIPTAAN\t: Program Komputer')
    
    doc.add_heading('DESKRIPSI SINGKAT', level=1)
    p = doc.add_paragraph(
        'Aplikasi Pasal Mobile adalah perangkat lunak berbasis mobile (Android/iOS) yang dirancang untuk memudahkan masyarakat dan profesional hukum dalam mengakses, mencari, dan mempelajari berbagai undang-undang dan peraturan di Indonesia. '
        'Aplikasi ini dilengkapi dengan fitur pencarian kata kunci, penyimpanan pasal favorit (bookmark/archive), dan mode baca yang nyaman. '
        'Selain itu, sistem ini didukung oleh Web Admin Dashboard untuk pengelolaan data undang-undang, pengguna, dan hak akses secara terpusat.'
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_page_break()
    
    doc.add_heading('DAFTAR FITUR DAN TAMPILAN', level=1)
    doc.add_paragraph('Berikut adalah deskripsi dari setiap tampilan dan fitur utama dalam aplikasi:')

    # Feature List
    features = [
        ("1. Tampilan Splash Screen", "Tampilan awal yang muncul saat aplikasi pertama kali dibuka, menampilkan logo aplikasi dan identitas visual \"Pasal\".", None),
        ("2. Tampilan Onboarding", "Halaman pengenalan fitur yang ditampilkan kepada pengguna baru untuk menjelaskan fungsi utama aplikasi sebelum masuk ke halaman login.", None),
        ("3. Halaman Login", "Form otentikasi pengguna untuk masuk ke dalam aplikasi. Pengguna harus memasukkan kredensial yang valid untuk mengakses fitur penuh.", None),
        ("4. Halaman Beranda (Home)", "Halaman utama yang menampilkan ringkasan, akses cepat ke undang-undang populer, atau kategori hukum. Berfungsi sebagai pusat navigasi aplikasi.", None),
        ("5. Halaman Perpustakaan (Library)", "Daftar lengkap undang-undang yang tersedia dalam database. Pengguna dapat menelusuri berbagai peraturan perundang-undangan di sini.", None),
        ("6. Halaman Detail Undang-Undang", "Menampilkan informasi rinci mengenai suatu undang-undang, termasuk nomor, tahun, tentang, dan daftar pasal-pasal di dalamnya.", "image36.jpeg"), # Guessing map
        ("7. Halaman Baca Pasal", "Tampilan mode baca untuk pasal tertentu. Dirancang agar nyaman dibaca dengan opsi navigasi antar pasal.", None),
        ("8. Fitur Pencarian (Search)", "Fitur untuk mencari undang-undang atau pasal spesifik berdasarkan kata kunci. Hasil pencarian akan menampilkan relevansi tertinggi.", "image32.jpeg"),
        ("9. Fitur Simpan (Bookmark/Archive)", "Memungkinkan pengguna untuk menyimpan pasal atau undang-undang tertentu ke dalam menu \"Tersimpan\" atau \"Arsip\" untuk akses cepat di kemudian hari.", "image33.jpeg"),
        ("10. Fitur Salin Teks", "Memudahkan pengguna untuk menyalin teks pasal untuk keperluan kutipan atau referensi dokumen lain.", "image34.jpeg"),
    ]

    used_images = set()

    for title, desc, img_name in features:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)
        if img_name:
            img_path = os.path.join(images_dir, img_name)
            if os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(3))
                    used_images.add(img_name)
                except Exception as e:
                    doc.add_paragraph(f"[Gagal memuat gambar: {img_name}]")
            else:
                 doc.add_paragraph("[Gambar tidak ditemukan, silakan tambahkan manual]")
        else:
             doc.add_paragraph("[Silakan tambahkan screenshot aplikasi mobile di sini]")

    doc.add_page_break()
    
    # Web Dashboard Section
    doc.add_heading('FITUR ADMIN DASHBOARD (WEB)', level=1)
    
    web_features = [
        ("11. Halaman Login Admin", "Portal masuk khusus untuk administrator guna mengelola konten dan pengguna aplikasi."),
        ("12. Dashboard Utama", "Menampilkan statistik ringkas mengenai jumlah pengguna, jumlah undang-undang, dan aktivitas sistem (Audit Log)."),
        ("13. Manajemen Pengguna (User Management)", "Fitur untuk melihat, menambah, mengedit, atau menonaktifkan akun pengguna aplikasi."),
        ("14. Manajemen Admin", "Fitur untuk mengelola akun administrator dan hak akses mereka."),
        ("15. Audit Log (Riwayat Aktivitas)", "Mencatat segala aktivitas yang terjadi dalam sistem untuk keperluan pemantauan dan keamanan."),
    ]

    for title, desc in web_features:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)
        doc.add_paragraph("[Silakan tambahkan screenshot Admin Dashboard di sini]")

    doc.add_page_break()
    doc.add_heading('LAMPIRAN GAMBAR TAMBAHAN', level=1)
    doc.add_paragraph('Berikut adalah gambar-gambar lain yang diekstrak dari laporan. Silakan pindahkan ke bagian yang sesuai di atas.')

    # Add all other images
    if os.path.exists(images_dir):
        all_images = sorted(os.listdir(images_dir))
        for img_file in all_images:
            if img_file not in used_images and (img_file.endswith('.png') or img_file.endswith('.jpeg') or img_file.endswith('.jpg')):
                doc.add_paragraph(f"Nama File: {img_file}")
                try:
                    doc.add_picture(os.path.join(images_dir, img_file), width=Inches(3))
                except:
                    pass
                doc.add_paragraph("-" * 20)

    doc.save(output_path)
    print(f"Document saved to: {output_path}")

if __name__ == "__main__":
    create_hki_document()
